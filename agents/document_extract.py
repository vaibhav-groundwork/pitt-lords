"""
Pure I/O utility for extracting plain text from uploaded lease files.

Deliberately contains no LLM calls, database access, or business logic.
Keeping extraction isolated here means:
  - It can be unit-tested with real files without standing up the full pipeline.
  - The extraction step can be swapped (e.g. switching PDF libraries) without
    touching any downstream agent or API code.
  - Failures in parsing surface as clear, typed errors before any expensive
    inference work begins.

Supported formats: .pdf (via pypdf), .docx (via python-docx).
For .docx files, both paragraph text and table cell text are extracted so
that rent schedules, late-fee tiers, and other tabular lease content are
not silently dropped.

The legacy binary .doc format is not supported -- callers should convert
to .docx before uploading.
"""

import sys
from pathlib import Path


def extract_text(file_path: str) -> str:
    """Extract and return all plain text from a lease document.

    Parameters
    ----------
    file_path:
        Absolute or relative path to the document. Must be a .pdf or .docx
        file. The legacy .doc binary format is explicitly rejected with a
        clear message rather than a cryptic library error.

    Returns
    -------
    str
        All text content joined as a single string. Pages (PDF) are separated
        by double newlines. For .docx files, paragraph text comes first,
        followed by a "--- Table content ---" section containing every table's
        cell text (cells joined with " | ", rows with newlines).

    Raises
    ------
    FileNotFoundError
        If ``file_path`` does not point to an existing file.
    ValueError
        If the file extension is not .pdf or .docx (including a specific
        message for .doc), or if the extracted text is empty or below the
        minimum character threshold after stripping whitespace. Silently
        returning empty text would be worse than raising here: every downstream
        legal-requirement check would produce a "not addressed" result that
        looks like a real finding rather than a parsing failure.
    """
    _MIN_CHARS = 50

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix == ".doc":
        raise ValueError(
            "The legacy .doc binary format is not supported. "
            "Please convert the file to .docx and re-upload."
        )
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            "Only .pdf and .docx files are accepted."
        )

    if len(text.strip()) < _MIN_CHARS:
        raise ValueError(
            "No extractable text found -- this may be a scanned image PDF "
            "without OCR, or an empty document."
        )

    return text


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF using pypdf.

    Each page's text is extracted individually and joined with a distinct
    ``[PAGE_BREAK]`` marker rather than a plain blank line. A plain ``\\n\\n``
    join was diagnosed as a real problem: agents/lease_parser.py's
    chunk_lease_text() treats every ``\\n\\n`` as a paragraph boundary, so a
    sentence that spans a PDF page break was split into two separate chunks,
    neither of which contained the complete sentence. A Fair Housing clause
    spanning a page break was confirmed affected: no downstream check could
    match it because the evidence only existed as the concatenation of two
    separate chunks.

    The ``[PAGE_BREAK]`` marker lets chunk_lease_text() (the matched
    counterpart to this join -- do not change one without the other) treat page
    breaks as a weaker signal than true blank-line paragraph boundaries, merging
    across the marker when the preceding fragment looks like a mid-sentence cut.
    """
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    # Matched pair with chunk_lease_text() in agents/lease_parser.py.
    # Use a named marker instead of a plain "\n\n" so the chunker can tell the
    # difference between a real paragraph break and a PDF rendering artefact.
    return "\n\n[PAGE_BREAK]\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    """Extract text from a Word document using python-docx.

    Paragraph text is extracted first, then all table content is appended
    under a "--- Table content ---" header. Exact document-order interleaving
    of paragraphs and tables is not attempted: the combined text is read as a
    whole by an LLM, so positional fidelity matters less than completeness.
    Tables are included because real leases frequently put substantive content
    (rent schedules, late-fee tiers, move-in charge breakdowns) in tabular
    form that would be silently dropped by a paragraph-only approach.

    Paragraphs are joined with double newlines (``\\n\\n``) so that
    agents/lease_parser.py's chunk_lease_text() -- which splits on ``\\n\\n``
    as its primary paragraph boundary -- treats each docx paragraph as a
    separate chunk. A single ``\\n`` join causes the entire document to
    collapse into one giant chunk, confirmed against the eval suite: every
    synthetic .docx test lease was producing exactly 1 chunk, causing most
    requirement_keys to never appear in any chunk's top-k retrieval results.

    Unlike the PDF case, no page-break-merging nuance is needed here. A docx
    ``Paragraph`` object is a clean, discrete, real paragraph by construction
    -- there is no rendering artefact equivalent to a PDF page break that might
    cut a sentence in half. A straightforward double-newline join is correct
    and sufficient.

    Table rows are also joined with double newlines for the same reason: each
    row is a discrete unit and should be able to become its own chunk in a
    large table, rather than all rows collapsing together into one.

    Text boxes are still not extracted -- python-docx does not expose them
    through the public API without XML traversal, and they are rarely used
    for substantive lease clauses.
    """
    from docx import Document

    doc = Document(str(path))

    paragraph_text = "\n\n".join(para.text for para in doc.paragraphs)

    table_rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            table_rows.append(row_text)

    if table_rows:
        table_block = "--- Table content ---\n" + "\n\n".join(table_rows)
        return paragraph_text + "\n\n" + table_block

    return paragraph_text


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python agents/document_extract.py <path-to-lease-file>")
        sys.exit(1)

    target = sys.argv[1]
    text = extract_text(target)
    char_count = len(text)

    print(f"Characters extracted: {char_count:,}")
    print("-" * 60)
    print(text[:2000])
